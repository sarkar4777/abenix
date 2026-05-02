"""Saudi Tourism Reports — auto-generated via Abenix st-report-generator agent."""
from __future__ import annotations

import io
import logging
import uuid

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.agent_utils import get_forge
from app.core.deps import get_db
from app.core.responses import error, success
from app.models.tourism_models import STReport, STDataset, STUser
from app.routers.auth import get_st_user

logger = logging.getLogger("sauditourism.reports")
router = APIRouter(prefix="/api/st/reports", tags=["st-reports"])


def _render_markdown(md: str) -> str:
    """Render Markdown to HTML using markdown-it-py with GFM tables + strikethrough.

    Falls back to <pre> escaped text if the lib is missing — preferable to the
    old regex hack that collapsed tables/lists to <br/>.
    """
    try:
        from markdown_it import MarkdownIt  # type: ignore

        md_renderer = MarkdownIt("commonmark", {"html": False, "linkify": True, "typographer": True})
        md_renderer.enable(["table", "strikethrough"])
        return md_renderer.render(md or "")
    except Exception as e:  # pragma: no cover
        logger.warning("markdown render failed: %s", e)
        from html import escape

        return f"<pre>{escape(md or '')}</pre>"


REPORT_TYPES = {
    "executive_briefing": {
        "title": "Executive Briefing",
        "description": "High-level summary of tourism performance with key metrics and trends",
        "agent_prompt": "Generate a comprehensive executive briefing for the KSA Ministry of Tourism. Include KPIs, regional performance, visitor trends, revenue analysis, and strategic recommendations. Use financial_calculator for all derived metrics.",
    },
    "regional_comparison": {
        "title": "Regional Comparison Report",
        "description": "Side-by-side comparison of all KSA tourism regions",
        "agent_prompt": "Create a detailed comparison of all KSA tourism regions including visitors, revenue, occupancy, and growth rates. Use financial_calculator to compute growth rates and rankings.",
    },
    "visitor_segmentation": {
        "title": "Visitor Segmentation Analysis",
        "description": "Breakdown by origin (domestic/GCC/international) and purpose (religious/leisure/business)",
        "agent_prompt": "Analyze visitor segmentation by origin (domestic/GCC/international) and purpose (religious/leisure/business). Include market share, growth trends, and spending patterns. Use financial_calculator for percentages and trends.",
    },
    "revenue_attribution": {
        "title": "Revenue Attribution Report",
        "description": "Where tourism money is spent: accommodation, F&B, transport, attractions, shopping",
        "agent_prompt": "Generate a revenue attribution analysis showing where tourism spending flows: accommodation, F&B, transport, attractions, shopping, etc. Use financial_calculator for sector shares and growth rates.",
    },
    "seasonal_analysis": {
        "title": "Seasonal Performance Analysis",
        "description": "Monthly/quarterly trends with seasonality patterns",
        "agent_prompt": "Analyze seasonal patterns in KSA tourism: monthly trends, peak/trough identification, Hajj/Umrah impact, seasonal pricing opportunities. Use financial_calculator for trend decomposition.",
    },
}


def _build_data_context(datasets: list[STDataset], max_chars: int = 50000) -> str:
    parts = []
    char_count = 0
    for d in datasets:
        if not d.raw_text:
            continue
        section = f"\n=== DATASET: {d.title} (type={d.dataset_type.value if d.dataset_type else 'general'}) ===\n"
        section += d.raw_text[:25000]
        if char_count + len(section) > max_chars:
            break
        parts.append(section)
        char_count += len(section)
    return "".join(parts) if parts else "NO DATA"


@router.get("/types")
async def get_report_types():
    # Strip agent_prompt from response
    clean = {k: {kk: vv for kk, vv in v.items() if kk != "agent_prompt"} for k, v in REPORT_TYPES.items()}
    return success(clean)


@router.post("/generate")
async def generate_report(
    body: dict,
    user: STUser = Depends(get_st_user),
    db: AsyncSession = Depends(get_db),
):
    report_type = body.get("type", "executive_briefing")
    if report_type not in REPORT_TYPES:
        return error(f"Unknown report type: {report_type}", 400)

    datasets = (await db.execute(
        select(STDataset).where(STDataset.user_id == user.id)
    )).scalars().all()

    if not datasets:
        return error("No datasets uploaded. Please upload tourism data first.", 400)

    preset = REPORT_TYPES[report_type]
    data_context = _build_data_context(datasets)
    forge, subject = get_forge(user)

    prompt = f"""{preset['agent_prompt']}

DATA FROM USER'S UPLOADED DATASETS:
{data_context}

FORMAT: Generate the report in rich Markdown format with:
- Executive summary at the top
- Tables for comparisons
- Bullet points for recommendations
- Bold key metrics
- Section headers for each topic

Use your tools (financial_calculator, structured_extractor, graph_builder) for all computations.
Do NOT make up data — use only what's in the datasets above."""

    result = await forge.execute("st-report-generator", prompt, act_as=subject)

    report = STReport(
        id=uuid.uuid4(),
        user_id=user.id,
        title=preset["title"],
        report_type=report_type,
        content=result.output,
        data={"agent_cost": result.cost, "agent_tokens": result.input_tokens + result.output_tokens},
    )
    db.add(report)
    await db.commit()

    return success({
        "id": str(report.id),
        "title": report.title,
        "type": report_type,
        "content": result.output,
        "content_html": _render_markdown(result.output),
        "agent_cost": result.cost,
        "created_at": report.created_at.isoformat() if report.created_at else None,
    })


@router.get("")
async def list_reports(
    user: STUser = Depends(get_st_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(STReport).where(STReport.user_id == user.id).order_by(STReport.created_at.desc())
    )
    reports = result.scalars().all()
    return success([
        {"id": str(r.id), "title": r.title, "type": r.report_type, "created_at": r.created_at.isoformat() if r.created_at else None}
        for r in reports
    ])


@router.get("/{report_id}")
async def get_report(
    report_id: str,
    user: STUser = Depends(get_st_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(STReport).where(STReport.id == uuid.UUID(report_id), STReport.user_id == user.id)
    )
    r = result.scalar_one_or_none()
    if not r:
        return error("Report not found", 404)
    return success({
        "id": str(r.id),
        "title": r.title,
        "type": r.report_type,
        "content": r.content,
        "content_html": _render_markdown(r.content or ""),
        "data": r.data,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    })


@router.post("/{report_id}/export")
async def export_report(
    report_id: str,
    format: str = Query("pdf", regex="^(pdf|docx)$"),
    user: STUser = Depends(get_st_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a report as PDF (weasyprint) or DOCX (python-docx)."""
    result = await db.execute(
        select(STReport).where(STReport.id == uuid.UUID(report_id), STReport.user_id == user.id)
    )
    r = result.scalar_one_or_none()
    if not r:
        return error("Report not found", 404)

    safe_title = (r.title or "report").replace(" ", "_")[:80]

    if format == "pdf":
        try:
            from weasyprint import HTML  # type: ignore
        except Exception as e:
            logger.error("weasyprint import failed: %s", e)
            return error("PDF export unavailable on this server", 500)

        html_body = _render_markdown(r.content or "")
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{r.title}</title>
<style>
  body {{ font-family: 'Liberation Sans', Helvetica, Arial, sans-serif; color:#0a2818; margin:32px; line-height:1.5; }}
  h1, h2, h3 {{ color:#16A34A; }}
  table {{ border-collapse: collapse; width:100%; margin:12px 0; }}
  th, td {{ border:1px solid #cfe7d8; padding:6px 10px; text-align:left; }}
  th {{ background:#e8f6ee; }}
  ul, ol {{ margin-left:18px; }}
  hr {{ border:0; border-top:1px solid #cfe7d8; margin:16px 0; }}
  code {{ background:#f1f5f4; padding:1px 4px; border-radius:3px; }}
</style></head><body>
<h1>{r.title}</h1>
{html_body}
</body></html>"""
        pdf_bytes = HTML(string=full_html).write_pdf()
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
        )

    # DOCX
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        logger.error("python-docx import failed: %s", e)
        return error("DOCX export unavailable on this server", 500)

    doc = Document()
    doc.add_heading(r.title or "Report", level=0)
    # Naive markdown -> docx: split on blank lines, treat # / ## / ### as headings,
    # | rows as table rows, - as bullets, else paragraph.
    lines = (r.content or "").splitlines()
    table_buf: list[list[str]] = []

    def _flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        cols = max(len(row) for row in table_buf)
        tbl = doc.add_table(rows=len(table_buf), cols=cols)
        for i, row in enumerate(table_buf):
            for j, cell in enumerate(row):
                tbl.rows[i].cells[j].text = cell
        table_buf = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("|") and line.count("|") >= 2:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            table_buf.append(cells)
            continue
        _flush_table()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    _flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )
