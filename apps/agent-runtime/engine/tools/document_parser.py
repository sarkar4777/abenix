"""Generic document text extraction tool."""

from __future__ import annotations

import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Any

from engine.tools.base import BaseTool, ToolResult

logger = logging.getLogger(__name__)

# Conservative size ceiling -- 500 MB should cover anything reasonable.
MAX_FILE_SIZE = 500 * 1024 * 1024

SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".csv": "csv",
    ".html": "html",
    ".htm": "html",
    ".md": "markdown",
    ".markdown": "markdown",
}

# Encoding fallback chain -- covers the vast majority of real-world files.
_ENCODING_CHAIN = ("utf-8", "latin-1", "cp1252")


class DocumentParserTool(BaseTool):
    name = "document_parser"
    description = (
        "Extract plain text from documents (PDF, DOCX, TXT, CSV, HTML, Markdown). "
        "Returns the full text content ready for analysis. Use as the first step "
        "before structured_extractor or any text analysis."
    )
    input_schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the file on the shared /data mount",
            },
            "max_chars": {
                "type": "integer",
                "description": "Maximum characters to return (default 100000)",
                "default": 100_000,
            },
            "page_range": {
                "type": "string",
                "description": (
                    "For PDFs only -- limit extraction to a page range, e.g. '1-5' "
                    "or '3' for a single page. Pages are 1-indexed."
                ),
            },
        },
        "required": ["file_path"],
    }

    # execute

    async def execute(self, arguments: dict[str, Any]) -> ToolResult:
        file_path = arguments.get("file_path", "")
        max_chars: int = int(arguments.get("max_chars", 100_000))
        page_range: str | None = arguments.get("page_range")

        if not file_path:
            return ToolResult(content="Error: file_path is required", is_error=True)

        path = Path(file_path)
        if not path.exists():
            return ToolResult(
                content=f"Error: File not found: {file_path}", is_error=True
            )

        if path.stat().st_size > MAX_FILE_SIZE:
            return ToolResult(
                content=f"Error: File exceeds {MAX_FILE_SIZE // (1024*1024)}MB limit",
                is_error=True,
            )

        ext = path.suffix.lower()
        file_type = SUPPORTED_EXTENSIONS.get(ext)
        if file_type is None:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS.keys()))
            return ToolResult(
                content=(
                    f"Error: Unsupported file type '{ext}'. "
                    f"Supported extensions: {supported}"
                ),
                is_error=True,
            )

        warnings: list[str] = []
        page_count: int | None = None

        try:
            if file_type == "pdf":
                text, page_count, pdf_warnings = self._read_pdf(path, page_range)
                warnings.extend(pdf_warnings)
            elif file_type == "docx":
                text, docx_warnings = self._read_docx(path)
                warnings.extend(docx_warnings)
            elif file_type == "csv":
                text = self._read_csv(path)
            elif file_type == "html":
                text = self._read_html(path)
            else:
                # txt, markdown -- plain text
                text = self._read_text(path)
        except Exception as exc:
            logger.exception("document_parser failed for %s", file_path)
            return ToolResult(
                content=f"Error: Failed to read file: {exc}",
                is_error=True,
            )

        # Enforce max_chars
        truncated = False
        total_chars = len(text)
        if total_chars > max_chars:
            text = text[:max_chars]
            truncated = True
            warnings.append(
                f"Output truncated from {total_chars} to {max_chars} characters"
            )

        result = {
            "text": text,
            "page_count": page_count,
            "char_count": len(text),
            "file_type": file_type,
            "warnings": warnings,
        }

        return ToolResult(
            content=json.dumps(result, default=str),
            metadata={
                "file_path": str(path),
                "file_type": file_type,
                "char_count": len(text),
                "page_count": page_count,
                "truncated": truncated,
            },
        )

    # PDF

    def _parse_page_range(
        self, page_range: str | None, total_pages: int
    ) -> tuple[int, int]:
        """Return (start_idx, end_idx) -- 0-based, end exclusive."""
        if not page_range:
            return 0, total_pages

        page_range = page_range.strip()
        try:
            if "-" in page_range:
                parts = page_range.split("-", 1)
                start = max(int(parts[0]) - 1, 0)
                end = min(int(parts[1]), total_pages)
            else:
                idx = int(page_range) - 1
                start = max(idx, 0)
                end = min(idx + 1, total_pages)
            return start, end
        except (ValueError, IndexError):
            return 0, total_pages

    def _read_pdf(
        self, path: Path, page_range: str | None
    ) -> tuple[str, int, list[str]]:
        warnings: list[str] = []
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            warnings.append("PyPDF2 is not installed -- PDF parsing unavailable")
            return "", 0, warnings

        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        start, end = self._parse_page_range(page_range, total_pages)

        pages: list[str] = []
        empty_pages = 0
        for i in range(start, end):
            page = reader.pages[i]
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text}")
            else:
                empty_pages += 1

        if empty_pages > 0:
            warnings.append(
                f"{empty_pages} page(s) returned no extractable text "
                "(may be scanned/OCR-only images)"
            )

        if not pages:
            warnings.append(
                "No text could be extracted from this PDF. "
                "The document may consist entirely of scanned images. "
                "Consider running OCR first."
            )
            return "", total_pages, warnings

        return "\n\n".join(pages), total_pages, warnings

    # DOCX

    def _read_docx(self, path: Path) -> tuple[str, list[str]]:
        warnings: list[str] = []
        try:
            from docx import Document
        except ImportError:
            return (
                "",
                [
                    "DOCX support requires python-docx. Install with: pip install python-docx"
                ],
            )

        doc = Document(str(path))
        parts: list[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Also extract text from tables (often important in business docs)
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        if not parts:
            warnings.append("No text content found in DOCX document")
            return "", warnings

        return "\n\n".join(parts), warnings

    # CSV

    def _read_csv(self, path: Path) -> str:
        raw = self._read_text_with_fallback(path)
        try:
            reader = csv.reader(io.StringIO(raw))
            rows = list(reader)
        except csv.Error:
            # Fall back to returning raw text
            return raw

        if not rows:
            return ""

        # Render as pipe-delimited text so downstream analysis is easy
        lines: list[str] = []
        header = rows[0]
        lines.append(" | ".join(header))
        lines.append("-" * max(len(lines[0]), 20))
        for row in rows[1:]:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    # HTML

    def _read_html(self, path: Path) -> str:
        raw = self._read_text_with_fallback(path)
        # Strip tags using regex (no external dep needed)
        # Remove script and style blocks entirely first
        text = re.sub(
            r"<(script|style)[^>]*>.*?</\1>", "", raw, flags=re.DOTALL | re.IGNORECASE
        )
        # Replace <br>, <p>, <div>, <li>, <tr> with newlines for readability
        text = re.sub(
            r"<(?:br|p|div|li|tr|h[1-6])[^>]*>", "\n", text, flags=re.IGNORECASE
        )
        # Strip remaining tags
        text = re.sub(r"<[^>]+>", "", text)
        # Decode common HTML entities
        text = (
            text.replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&quot;", '"')
            .replace("&#39;", "'")
            .replace("&nbsp;", " ")
        )
        # Collapse excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # Plain text (TXT / Markdown)

    def _read_text(self, path: Path) -> str:
        return self._read_text_with_fallback(path)

    # Helpers

    def _read_text_with_fallback(self, path: Path) -> str:
        """Try the encoding chain until one succeeds."""
        for encoding in _ENCODING_CHAIN:
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, ValueError):
                continue
        # Last resort: read bytes and decode lossily
        return path.read_bytes().decode("utf-8", errors="replace")
