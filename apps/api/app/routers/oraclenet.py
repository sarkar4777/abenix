"""OracleNet — Strategic Decision Analysis API."""

from __future__ import annotations

import json
import sys
import uuid
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Body, Depends, Query
from fastapi.responses import JSONResponse, StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.deps import get_current_user, get_db
from app.core.responses import error, success
from app.schemas.oraclenet import AnalyzeRequest

sys.path.insert(0, str(Path(__file__).resolve().parents[4] / "packages" / "db"))

from models.agent import Agent
from models.execution import Execution
from models.user import User

router = APIRouter(prefix="/api/oraclenet", tags=["oraclenet"])


@router.post("/analyze")
async def analyze_decision(
    body: AnalyzeRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> Any:
    """Submit a decision for OracleNet analysis. Returns SSE stream of agent progress."""
    # Check per-user quota
    from app.core.usage import check_user_quota

    quota_error = await check_user_quota(user)
    if quota_error:
        return error(quota_error, 429)

    # Look up the OracleNet pipeline agent
    result = await db.execute(
        select(Agent).where(
            Agent.slug == "oraclenet-pipeline", Agent.status == "active"
        )
    )
    agent = result.scalar_one_or_none()
    if not agent:
        return error(
            "OracleNet pipeline agent not found. Run seed_agents.py first.", 404
        )

    model_cfg = agent.model_config_ or {}
    pipeline_config = model_cfg.get("pipeline_config")
    if not pipeline_config or not pipeline_config.get("nodes"):
        return error("OracleNet pipeline agent has no pipeline configuration", 400)

    tool_names = model_cfg.get("tools", [])

    # Build context with the decision prompt.
    # `depth` is consumed by the depth_router node in oraclenet_pipeline.yaml
    # to decide which optional agents to run (quick=3, standard=5, deep=7).
    context = {
        "message": body.decision_prompt,
        "depth": body.depth or "standard",
        **(body.context or {}),
    }

    # Create execution record
    from models.execution import ExecutionStatus

    execution = Execution(
        tenant_id=user.tenant_id,
        agent_id=agent.id,
        user_id=user.id,
        input_message=body.decision_prompt,
        status=ExecutionStatus.RUNNING,
        model_used="pipeline",
    )
    db.add(execution)
    await db.commit()
    await db.refresh(execution)

    # Reuse the pipeline streaming helper from the agents router
    from app.routers.agents import _stream_pipeline_execution

    return StreamingResponse(
        _stream_pipeline_execution(
            execution_id=execution.id,
            message=body.decision_prompt,
            tool_names=tool_names,
            pipeline_config=pipeline_config,
            db=db,
            context=context,
            tenant_id=str(user.tenant_id),
            agent_id=str(agent.id),
            agent_name=agent.name,
            timeout_seconds=600,  # OracleNet: 7 agents with web research need up to 10 min
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.get("/sessions")
async def list_sessions(
    limit: int = Query(default=20, le=100),
    offset: int = Query(default=0),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    """List past OracleNet decision analyses."""
    # Find the OracleNet pipeline agent
    result = await db.execute(
        select(Agent.id).where(Agent.slug == "oraclenet-pipeline")
    )
    agent_row = result.first()
    if not agent_row:
        return success([])

    agent_id = agent_row[0]

    result = await db.execute(
        select(Execution)
        .where(Execution.agent_id == agent_id, Execution.user_id == user.id)
        .order_by(Execution.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    executions = result.scalars().all()

    items = [
        {
            "execution_id": str(e.id),
            "decision_prompt": (e.input_message or "")[:200],
            "status": e.status.value if e.status else "unknown",
            "confidence": float(e.confidence_score) if e.confidence_score else None,
            "created_at": e.created_at.isoformat() if e.created_at else None,
            "duration_ms": e.duration_ms,
            "cost": float(e.cost) if e.cost else None,
        }
        for e in executions
    ]
    return success(items)


@router.get("/sessions/{execution_id}")
async def get_session(
    execution_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    """Get a completed OracleNet Decision Brief."""
    result = await db.execute(
        select(Execution).where(
            Execution.id == execution_id,
            Execution.user_id == user.id,
        )
    )
    execution = result.scalar_one_or_none()
    if not execution:
        return error("Session not found", 404)

    # Parse the synthesizer output as Decision Brief
    brief = None
    if execution.output_message:
        try:
            brief = json.loads(execution.output_message)
        except (json.JSONDecodeError, TypeError):
            brief = {"raw_output": execution.output_message}

    return success(
        {
            "execution_id": str(execution.id),
            "status": execution.status.value if execution.status else "unknown",
            "decision_prompt": execution.input_message,
            "brief": brief,
            "node_results": execution.node_results,
            "duration_ms": execution.duration_ms,
            "cost": float(execution.cost) if execution.cost else None,
            "tokens": {
                "input": execution.input_tokens,
                "output": execution.output_tokens,
            },
            "created_at": (
                execution.created_at.isoformat() if execution.created_at else None
            ),
        }
    )


def _brief_to_markdown(brief: dict[str, Any]) -> str:
    """Render a Decision Brief dict as readable markdown."""
    lines: list[str] = ["# OracleNet Decision Brief", ""]
    if brief.get("summary") or brief.get("executive_summary"):
        lines += [
            "## Executive Summary",
            str(brief.get("summary") or brief.get("executive_summary") or ""),
            "",
        ]
    if brief.get("confidence") is not None:
        lines += [f"**Confidence:** {brief.get('confidence')}%", ""]
    if brief.get("recommendation"):
        rec = brief["recommendation"]
        rec_text = (
            rec if isinstance(rec, str) else (rec.get("action") or json.dumps(rec))
        )
        lines += ["## Recommendation", str(rec_text), ""]
    for s in brief.get("scenarios", []) or []:
        title = s.get("title") or s.get("name") or "Scenario"
        prob = s.get("probability") or 0
        lines += [
            f"### Scenario: {title} ({prob}%)",
            str(s.get("description") or ""),
            "",
        ]
    if brief.get("stakeholders"):
        lines += ["## Stakeholders"]
        for s in brief["stakeholders"]:
            lines += [
                f"- **{s.get('name', 'Stakeholder')}** [{s.get('sentiment', 'neutral')}]: {s.get('details') or s.get('impact') or ''}"
            ]
        lines += [""]
    if brief.get("risks"):
        lines += ["## Risks"]
        for r in brief["risks"]:
            lines += [
                f"- **{r.get('title', 'Risk')}** [{r.get('severity', 'medium')}] {r.get('probability', 0)}%: {r.get('description') or ''}"
            ]
        lines += [""]
    if brief.get("conditions"):
        lines += ["## Conditions"] + [f"- {c}" for c in brief["conditions"]] + [""]
    if brief.get("monitoringTriggers") or brief.get("monitoring_triggers"):
        triggers = (
            brief.get("monitoringTriggers") or brief.get("monitoring_triggers") or []
        )
        lines += ["## Monitoring Triggers"] + [f"- {t}" for t in triggers] + [""]
    return "\n".join(lines)


@router.post("/export")
async def export_brief_inline(
    format: str = Query(
        default="pdf", description="Output format: pdf | docx | markdown"
    ),
    payload: dict[str, Any] = Body(...),
    user: User = Depends(get_current_user),
) -> Any:
    """Export a Decision Brief supplied inline as PDF or DOCX.

    Body shape: {"brief": {... BriefData ...}}.
    """
    fmt = (format or "").lower()
    if fmt not in ("pdf", "docx", "markdown", "md"):
        return error("format must be pdf, docx, or markdown", 400)

    brief = payload.get("brief") or payload
    if not isinstance(brief, dict):
        return error("Body must include {brief: {...}}", 400)

    md_text = _brief_to_markdown(brief)

    if fmt in ("markdown", "md"):
        return StreamingResponse(
            iter([md_text.encode()]),
            media_type="text/markdown",
            headers={
                "Content-Disposition": 'attachment; filename="oraclenet-brief.md"'
            },
        )

    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            import io

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            for raw_para in md_text.split("\n\n"):
                para = raw_para.strip()
                if not para:
                    continue
                if para.startswith("# "):
                    story.append(Paragraph(para.lstrip("# ").strip(), styles["Title"]))
                elif para.startswith("## "):
                    story.append(
                        Paragraph(para.lstrip("# ").strip(), styles["Heading2"])
                    )
                elif para.startswith("### "):
                    story.append(
                        Paragraph(para.lstrip("# ").strip(), styles["Heading3"])
                    )
                else:
                    # Escape <, > so reportlab doesn't try to parse them as markup
                    safe = (
                        para.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(safe, styles["Normal"]))
                story.append(Spacer(1, 6))
            doc.build(story)
            buffer.seek(0)
            return StreamingResponse(
                iter([buffer.read()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": 'attachment; filename="oraclenet-brief.pdf"'
                },
            )
        except ImportError:
            return error("reportlab not installed for PDF generation", 500)
        except Exception as e:
            return error(f"PDF generation failed: {e}", 500)

    if fmt == "docx":
        try:
            from docx import Document
            import io

            doc = Document()
            for raw_para in md_text.split("\n\n"):
                para = raw_para.strip()
                if not para:
                    continue
                if para.startswith("# "):
                    doc.add_heading(para.lstrip("# ").strip(), 0)
                elif para.startswith("## "):
                    doc.add_heading(para.lstrip("# ").strip(), level=1)
                elif para.startswith("### "):
                    doc.add_heading(para.lstrip("# ").strip(), level=2)
                else:
                    doc.add_paragraph(para)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return StreamingResponse(
                iter([buffer.read()]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": 'attachment; filename="oraclenet-brief.docx"'
                },
            )
        except ImportError:
            return error("python-docx not installed for DOCX generation", 500)
        except Exception as e:
            return error(f"DOCX generation failed: {e}", 500)

    return error("Unsupported format", 400)


@router.get("/sessions/{execution_id}/export/{fmt}")
async def export_brief(
    execution_id: uuid.UUID,
    fmt: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> Any:
    """Export Decision Brief as PDF, DOCX, or Markdown."""
    if fmt not in ("pdf", "docx", "markdown", "md"):
        return error("Format must be: pdf, docx, markdown", 400)

    result = await db.execute(
        select(Execution).where(
            Execution.id == execution_id,
            Execution.user_id == user.id,
        )
    )
    execution = result.scalar_one_or_none()
    if not execution:
        return error("Session not found", 404)

    brief_text = execution.output_message or ""

    if fmt in ("markdown", "md"):
        return StreamingResponse(
            iter([brief_text.encode()]),
            media_type="text/markdown",
            headers={
                "Content-Disposition": f'attachment; filename="decision-brief-{execution_id}.md"'
            },
        )

    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            import io

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            story.append(Paragraph("OracleNet Decision Brief", styles["Title"]))
            story.append(Spacer(1, 12))

            # Split brief into paragraphs
            for para in brief_text.split("\n\n"):
                if para.strip():
                    if para.strip().startswith("#"):
                        clean = para.strip().lstrip("#").strip()
                        story.append(Paragraph(clean, styles["Heading2"]))
                    else:
                        story.append(Paragraph(para.strip(), styles["Normal"]))
                    story.append(Spacer(1, 6))

            doc.build(story)
            buffer.seek(0)
            return StreamingResponse(
                iter([buffer.read()]),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="decision-brief-{execution_id}.pdf"'
                },
            )
        except ImportError:
            return error("reportlab not installed for PDF generation", 500)

    if fmt == "docx":
        try:
            from docx import Document
            import io

            doc = Document()
            doc.add_heading("OracleNet Decision Brief", 0)

            for para in brief_text.split("\n\n"):
                if para.strip():
                    if para.strip().startswith("#"):
                        clean = para.strip().lstrip("#").strip()
                        doc.add_heading(clean, level=2)
                    else:
                        doc.add_paragraph(para.strip())

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return StreamingResponse(
                iter([buffer.read()]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="decision-brief-{execution_id}.docx"'
                },
            )
        except ImportError:
            return error("python-docx not installed for DOCX generation", 500)
